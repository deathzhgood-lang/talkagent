from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "TalkAgent人工智能大作业报告.docx"
ASSETS = ROOT / "tmp" / "report_assets"


FONT_CN = "宋体"
FONT_EN = "Times New Roman"
HEADING_COLOR = RGBColor(46, 116, 181)
MUTED_COLOR = RGBColor(90, 90, 90)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_COLOR, 16, 8),
        ("Heading 2", 13, HEADING_COLOR, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("TalkAgent 本地 AI 客服系统设计与实现报告")
    set_run_font(r, 9, False, MUTED_COLOR)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("TalkAgent 本地 AI 客服系统设计与实现报告")
    set_run_font(r, 18, True, RGBColor(0, 0, 0))


def para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, 11)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, 11)
    return p


def number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, 11)
    return p


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, fill: str | None = None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 10, bold)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(tbl.rows[0].cells[i], h, True, "F2F4F7")
        tbl.rows[0].cells[i].width = Inches(widths[i])
    for row in rows:
        cells = tbl.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return tbl


def caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, 9, False, MUTED_COLOR)


def code_block(doc: Document, code: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, 10.5, True, RGBColor(31, 77, 120))

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, "F6F8FA")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(code.strip("\n").splitlines()):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_picture(doc: Document, image_name: str, caption_text: str, width_in: float = 6.1):
    image_path = ASSETS / image_name
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(width_in))
        caption(doc, caption_text)


def page_break(doc: Document):
    doc.add_page_break()


def build():
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("摘要", level=1)
    para(
        doc,
        "本项目设计并实现了一个名为 TalkAgent 的本地 AI 客服系统。系统以检索增强生成（RAG）为核心思路，支持用户上传 PDF、Word、Markdown、文本和图片资料，将资料解析、分块、向量化后存入本地知识库。用户提问时，系统先判断问题是否需要查阅文档，再检索相关片段并调用本地大模型生成回答，同时展示答案来源。项目重点验证了“上传时生成知识目录，提问时先由智能体判断是否查文档，再结合文档和通用知识回答，最后进行自查”的方案可行性。测试结果表明，TalkAgent 已经完成文档入库、检索问答、来源追溯和本地桌面交互的基本闭环，可作为一个小型本地 AI 客服原型。"
    )
    p = doc.add_paragraph()
    r = p.add_run("关键词：")
    set_run_font(r, 11, True)
    r = p.add_run("人工智能；RAG；本地大模型；向量数据库；智能客服；文档问答")
    set_run_font(r, 11)
    page_break(doc)

    doc.add_heading("1 项目概述", level=1)
    para(
        doc,
        "随着大语言模型的发展，企业和个人希望把已有文档快速转化为可对话的知识库。传统客服系统往往依赖人工维护问答规则，更新成本较高；通用大模型虽然具备较强语言能力，但如果不了解用户上传的私有资料，就容易回答不准确或编造内容。因此，本项目选择构建一个本地 RAG 智能客服系统，使模型在回答前先检索用户资料，再根据检索结果组织答案。"
    )
    para(
        doc,
        "TalkAgent 的定位是一个可在 Windows 本地运行的 AI 客服原型。用户通过桌面界面上传资料、刷新知识库、进行对话并查看来源。系统不要求用户打开浏览器，也不依赖云端模型完成核心回答流程，适合课程实验、个人知识库和小型业务客服场景。"
    )
    doc.add_heading("1.1 研究背景", level=2)
    para(
        doc,
        "在实际业务中，客服人员经常需要查阅大量产品说明、接口文档、岗位手册和常见问题资料。人工查找资料速度慢，且不同人员对同一文档的理解可能不一致。大语言模型虽然能生成流畅回答，但如果直接让模型回答业务问题，它并不知道本地私有文档中的具体规则，容易出现“看起来合理但没有依据”的回答。"
    )
    para(
        doc,
        "RAG 技术正好适合解决这一类问题。它不是把全部资料重新训练进模型，而是在提问时从资料库中检索相关内容，再让模型基于这些内容回答。这样既能利用大模型的语言理解和总结能力，又能让回答尽量贴近用户上传的资料。"
    )
    doc.add_heading("1.2 项目意义", level=2)
    for item in [
        "对课程学习而言，本项目把大语言模型、向量检索、文档解析、桌面应用开发等知识结合在一个完整系统中，能够体现人工智能应用开发的综合能力。",
        "对实际应用而言，本项目展示了一个轻量级本地知识库客服的基本形态，适合在不方便上传资料到云端的场景中使用。",
        "对后续扩展而言，项目预留了 OCR、知识目录、问题路由和答案自查等能力，可以继续发展为更完整的企业内部知识助手。",
    ]:
        bullet(doc, item)

    doc.add_heading("2 需求分析", level=1)
    doc.add_heading("2.1 功能需求", level=2)
    for item in [
        "文档上传：支持 PDF、DOCX、TXT、MD、PNG、JPG、JPEG 等常见资料格式。",
        "知识库构建：能够自动解析文档内容，进行文本分块、向量化并保存到本地向量数据库。",
        "智能问答：用户输入自然语言问题后，系统能够根据文档内容和必要的通用知识生成回答。",
        "来源展示：系统需要展示回答所依据的文件名、片段编号和原文摘要，便于用户核查。",
        "本地运行：界面以桌面应用形式提供，不要求用户通过浏览器访问。",
        "降级可用：当本地大模型或完整 embedding 模型暂时不可用时，系统仍能返回检索片段摘要，避免整体不可用。",
    ]:
        bullet(doc, item)

    doc.add_heading("2.2 非功能需求", level=2)
    para(
        doc,
        "系统需要尽量保持轻量、易部署、易验证。由于实验环境主要依赖普通 CPU，本项目选择较小规模的本地模型，并通过降级策略保证流程可运行。同时，回答过程要尽量避免无依据编造，对缺少资料的问题给出明确提示。"
    )
    doc.add_heading("2.3 使用对象与输入输出", level=2)
    para(
        doc,
        "系统的主要使用对象是需要频繁查阅资料的普通用户，例如课程项目演示者、企业内部客服人员、接口文档使用者或岗位资料学习者。用户不需要了解向量数据库和大模型调用细节，只需要通过界面上传文档并提出自然语言问题。"
    )
    table(
        doc,
        ["项目", "内容"],
        [
            ["输入资料", "PDF、Word、Markdown、TXT、图片等形式的业务资料或学习资料"],
            ["用户问题", "自然语言问题，例如“某接口有哪些必填参数”“应聘 AI 应用开发需要哪些能力”"],
            ["系统输出", "带有回答标签的自然语言答复，以及相关文档来源片段"],
            ["异常输出", "当模型不可用或文档缺少依据时，给出不可用原因、检索摘要或补充资料建议"],
        ],
        [1.5, 4.8],
    )
    doc.add_heading("2.4 评价标准", level=2)
    for item in [
        "可用性：系统能否完成上传、入库、检索、回答、来源展示的完整流程。",
        "准确性：回答是否尽量依据文档，是否减少无依据编造。",
        "完整性：对于参数列表、岗位要求等枚举类问题，是否尽量避免漏项。",
        "可解释性：用户是否能看到回答所依据的文件和片段。",
        "鲁棒性：模型、OCR 或 embedding 不完整时，系统是否仍有可理解的降级结果。",
    ]:
        bullet(doc, item)
    page_break(doc)

    doc.add_heading("3 核心原理与技术栈", level=1)
    doc.add_heading("3.1 RAG 原理", level=2)
    para(
        doc,
        "RAG 是 Retrieval-Augmented Generation 的缩写，中文通常称为“检索增强生成”。它的基本思想是：当用户提出问题时，系统先从知识库中检索相关资料，再把这些资料作为上下文交给大语言模型生成回答。这样模型不需要重新训练，也可以利用用户自己的文档知识。"
    )
    for item in [
        "第一步，文档解析：把用户上传的资料转换为纯文本。",
        "第二步，文本分块：把长文档切成较短片段，避免一次检索内容过大。",
        "第三步，向量化：用 embedding 模型把文本片段转换成向量。",
        "第四步，相似度检索：把用户问题也转换成向量，找出最相近的文档片段。",
        "第五步，生成回答：把问题、检索片段和历史对话交给大模型，生成最终答复。",
    ]:
        number(doc, item)

    doc.add_heading("3.2 智能体流程设计", level=2)
    para(
        doc,
        "在普通 RAG 的基础上，本项目进一步加入了智能体式的判断流程。核心想法是：上传文档时先生成 Markdown 总结和检索目录；用户提问时，先把问题和目录交给模型或路由模块判断是否需要查文档；如果需要，再检索相关文档片段；回答生成后，再进行一次遗漏和事实自查。该流程可以减少盲目检索，也能让系统同时处理通用问题、文档问题和综合问题。"
    )
    table(
        doc,
        ["阶段", "处理内容", "作用"],
        [
            ["上传阶段", "生成文档摘要和知识目录", "帮助系统了解知识库中有哪些资料"],
            ["路由阶段", "判断 general、document、hybrid、missing_evidence", "决定是否需要查文档以及如何回答"],
            ["检索阶段", "按问题扩展查询并检索相关片段", "提高召回率，减少漏答"],
            ["生成阶段", "结合文档、上下文和通用知识回答", "形成可读的客服答复"],
            ["自查阶段", "检查必填项、事实依据和遗漏", "降低关键内容漏掉的概率"],
        ],
        [1.0, 2.8, 2.5],
    )

    doc.add_heading("3.3 技术栈", level=2)
    table(
        doc,
        ["模块", "技术", "说明"],
        [
            ["界面层", "Tkinter", "构建 Windows 本地桌面应用，提供上传、对话和来源查看功能"],
            ["大模型", "Ollama + codex-app", "用于本地生成自然语言回答"],
            ["向量模型", "BGE-M3 / hash embedding", "BGE-M3 用于语义向量化，权重不完整时使用 hash 方案兜底"],
            ["向量数据库", "ChromaDB", "持久化保存文档片段和向量索引"],
            ["文档解析", "PyMuPDF、docx2txt、OCR/视觉模型", "解析 PDF、Word、文本和图片资料"],
            ["后端逻辑", "Python", "完成文档处理、问答链路、会话记录和配置管理"],
        ],
        [1.1, 1.9, 3.2],
    )
    doc.add_heading("3.4 Embedding 与向量检索说明", level=2)
    para(
        doc,
        "Embedding 可以理解为把一段文本转换成一组数字。语义相近的文本，其向量距离也会比较接近。例如“必填参数有哪些”和“required fields”虽然表达不同，但在语义上都与参数要求相关，向量检索可以帮助系统找到相近片段。TalkAgent 默认使用 BGE-M3 作为 embedding 模型；在权重不完整时，系统使用 hash embedding 兜底，保证流程可运行，但语义效果会弱于真实模型。"
    )
    para(
        doc,
        "向量数据库 ChromaDB 的作用是保存这些文本片段和对应向量。当用户提问时，系统把问题也转换成向量，并在数据库中寻找相似度最高的片段。相比简单关键词搜索，向量检索更适合处理同义表达和自然语言问题。"
    )
    doc.add_heading("3.5 问题路由与答案自查说明", level=2)
    para(
        doc,
        "问题路由是本项目相对普通 RAG 的重要改进。它的作用是先判断用户问题属于哪种类型：如果是普通概念问题，就使用通用知识回答；如果明显需要文档依据，就进入文档检索；如果需要同时结合文档和常识，就进入综合回答；如果文档缺少依据，则提示缺少资料。"
    )
    para(
        doc,
        "答案自查主要用于解决“回答漏项”问题。例如用户询问接口必填参数时，模型可能只回答部分字段。为降低这种风险，系统在检索阶段会扩展查询词，在生成后还会抽取文档中标注为“必选、必填、required”的参数，对比回答是否遗漏。"
    )
    page_break(doc)

    doc.add_heading("4 系统总体设计", level=1)
    para(
        doc,
        "TalkAgent 采用分层设计。最上层是桌面交互界面，负责接收用户上传和提问；中间层是应用逻辑，包括文档加载、文本分块、知识目录、问题路由、问答链和答案自查；底层是本地模型服务、向量数据库和文件存储。这样的设计使界面、检索、生成和存储之间相对独立，方便后续调试和扩展。"
    )
    table(
        doc,
        ["层次", "主要文件", "功能"],
        [
            ["界面层", "app/desktop.py", "桌面 UI、上传面板、聊天窗口、来源显示"],
            ["文档处理层", "document_loader.py、text_splitter.py、ocr.py", "文档解析、图片识别、中文分块"],
            ["知识库层", "vector_store.py、knowledge_index.py", "向量入库、检索、文档摘要目录"],
            ["智能问答层", "question_router.py、rag_chain.py、answer_checker.py", "问题路由、检索增强回答、答案自查"],
            ["配置与记录", "config.py、chat_store.py", "环境配置、对话历史持久化"],
        ],
        [1.1, 2.4, 2.8],
    )
    doc.add_heading("4.1 数据流设计", level=2)
    for item in [
        "用户通过桌面界面选择文件，系统读取文件字节并保存到 data/uploads 目录。",
        "文档加载模块根据文件类型解析文本，图片内容尝试通过 OCR 或视觉模型转换为文字。",
        "文本分块模块按中文友好的规则切分文档，尽量避免把同一语义字段切断。",
        "向量库模块将片段写入 ChromaDB，并记录 file_id、file_name、chunk_index 等元数据。",
        "知识目录模块为每个文件生成 Markdown 摘要，并维护全局 index.json，供问题路由使用。",
        "用户提问后，问答链根据路由结果检索片段，调用本地模型生成回答，并把来源返回到界面。",
    ]:
        number(doc, item)
    doc.add_heading("4.2 界面设计", level=2)
    para(
        doc,
        "界面设计遵循“以对话为主、上传为辅”的原则。主区域用于显示用户和助手的对话，双方消息有明显区分；上传区域通过“文档”按钮展开，避免长期占用界面空间；来源区域可以折叠查看，既不干扰阅读，又能在需要时核查依据。"
    )
    para(
        doc,
        "相比浏览器页面，本地桌面界面对普通用户更直观，也更符合“本地部署程序”的要求。用户只需要双击启动脚本即可进入应用，后续操作集中在同一个窗口内完成。"
    )
    doc.add_heading("4.3 存储设计", level=2)
    table(
        doc,
        ["目录", "保存内容", "作用"],
        [
            ["data/uploads", "原始上传文件", "便于后续重新解析和管理"],
            ["data/chroma", "向量数据库持久化文件", "保存可检索的文档片段向量"],
            ["data/conversations", "对话历史 JSON", "支持多轮对话和历史记录"],
            ["data/knowledge_index", "文档摘要与检索目录", "辅助智能体判断是否查文档"],
            ["models", "本地模型缓存", "保存 embedding 等模型文件"],
        ],
        [1.7, 2.1, 2.5],
    )
    page_break(doc)

    doc.add_heading("5 具体实现步骤", level=1)
    for item in [
        "搭建项目骨架：建立 app 目录，补齐配置、文档加载、向量库、问答链、聊天记录、桌面界面等模块。",
        "实现配置集中管理：通过 .env 读取模型地址、模型名称、分块大小、检索数量、数据目录等参数。",
        "实现文档上传与解析：用户上传文件后，系统保存原始文件并提取文本；对于图片或含图片文档，尝试通过 OCR 或视觉模型解析文字内容。",
        "实现文本分块与向量入库：将文档切分成适合检索的片段，调用 embedding 模型生成向量，并写入 ChromaDB。",
        "实现知识目录生成：上传文档后生成 Markdown 摘要和全局目录，用于后续问题路由。",
        "实现问题路由：根据用户问题和知识目录判断问题属于通用问题、文档问题、综合问题还是缺少依据问题。",
        "实现 RAG 回答：对需要文档的问题进行相似度检索，将相关片段、文档摘要、历史对话和用户问题组织成 prompt，交给本地模型回答。",
        "实现答案自查：对参数枚举、必填项等容易漏答的问题，进行二次检查和确定性补全。",
        "实现桌面交互：将上传、删除、刷新、对话、来源展示等功能整合到简洁的 Tkinter 桌面应用。",
        "实现降级策略：当 Ollama 不可用时，返回检索片段摘要；当 BGE-M3 权重缺失时，使用 hash embedding 保证流程可继续验证。",
    ]:
        number(doc, item)
    doc.add_heading("5.1 文档上传与入库实现", level=2)
    para(
        doc,
        "文档上传流程从桌面界面开始。用户点击上传按钮后，系统读取文件内容，并根据扩展名选择对应解析方式。普通文本和 Markdown 可以直接读取；PDF 使用 PyMuPDF 提取文字；Word 文档使用 docx2txt；图片类文件则尝试进入 OCR 或视觉模型处理流程。解析后的文本会被统一转换为 Document 对象，便于后续分块和入库。"
    )
    para(
        doc,
        "文本分块时，本项目使用 chunk_size=500、chunk_overlap=100 的配置。这样可以让每个片段长度适中，同时通过重叠部分保留上下文，减少关键信息刚好被切开的情况。每个片段入库时都带有文件名、文件 ID、片段编号等元数据，后续来源展示就依赖这些信息。"
    )
    doc.add_heading("5.2 问答链实现", level=2)
    para(
        doc,
        "问答链的核心文件是 rag_chain.py。该模块先读取最近几轮对话作为历史上下文，再调用问题路由模块判断当前问题类型。对于需要文档的问题，系统会调用 similarity_search 检索相关片段，并把文档摘要、检索片段、历史对话和用户问题一起组织成 prompt，发送给本地模型。"
    )
    para(
        doc,
        "为了适应参数列表类问题，系统加入了扩展查询策略。例如问题中包含“必选、必填、参数、required”等关键词时，会额外生成与 required、prompt、Action、Version 等相关的查询，提高召回更多参数片段的概率。"
    )
    doc.add_heading("5.3 降级策略实现", level=2)
    para(
        doc,
        "本地应用经常会遇到模型服务未启动、模型未下载、网络不可用等问题。如果没有降级处理，用户体验会变成直接报错。本项目在 LLM 调用失败时，检查是否已经检索到相关文档片段；如果有，就返回片段摘要和不可用原因；如果没有，则明确提示本地大模型不可用且没有可参考资料。"
    )
    para(
        doc,
        "Embedding 也设置了类似思路。完整 BGE-M3 权重缺失时，系统使用 hash embedding 兜底。虽然 hash embedding 的语义效果有限，但它能保证文档入库、检索流程和界面交互继续运行，便于项目演示和后续调试。"
    )
    page_break(doc)

    doc.add_heading("6 测试结果与分析", level=1)
    para(
        doc,
        "本项目主要从文档入库、问答效果、来源追溯和异常降级四个方面进行测试。测试样例包括 AI 应用开发岗位资料、即梦 AI 图片生成接口文档等。系统能够完成上传、入库、检索、回答和展示来源的基本闭环。"
    )
    doc.add_heading("6.1 测试环境", level=2)
    table(
        doc,
        ["项目", "配置或说明"],
        [
            ["运行系统", "Windows 本地环境"],
            ["运行方式", "Tkinter 桌面应用"],
            ["大模型服务", "Ollama，本地模型名称 codex-app"],
            ["向量数据库", "ChromaDB 本地持久化"],
            ["测试资料", "AI 应用开发岗位资料、RAG 测试文档、即梦 AI 图片生成接口文档"],
            ["测试目标", "验证上传、入库、检索、回答、来源展示和异常降级"],
        ],
        [1.8, 4.6],
    )
    doc.add_heading("6.2 测试用例", level=2)
    table(
        doc,
        ["测试项", "测试内容", "结果"],
        [
            ["文档入库", "上传 4 个测试文件并查看知识库统计", "成功入库，共生成 122 个片段"],
            ["文档问答", "询问“应聘 AI 应用开发需要怎样才能被录用”", "系统基于岗位文档回答技术能力、产品思维、伦理合规和实践经验"],
            ["来源追溯", "展开来源面板查看命中文档片段", "可以看到文件名、片段编号和原文摘要"],
            ["必填参数问题", "询问即梦 4.6 接口必选参数", "通过多查询检索和自查机制减少漏答 prompt 等关键字段"],
            ["模型不可用", "Ollama 服务异常或返回 502", "系统返回检索片段摘要和不可用原因，不直接崩溃"],
        ],
        [1.2, 3.0, 2.0],
    )
    doc.add_heading("6.3 测试结果分析", level=2)
    para(
        doc,
        "从测试结果看，系统已经具备一个本地 AI 客服的基础闭环。知识库统计能够显示文件数量和片段数量，说明上传、解析、分块和入库流程可用；对岗位能力问题的回答能够引用文档内容，说明检索和生成链路有效；来源面板能够展示文件名和片段摘要，说明回答具有一定可解释性。"
    )
    para(
        doc,
        "同时，测试也暴露出系统质量依赖底层组件的问题。例如图片内容识别依赖 OCR 或视觉模型，如果 OCR 未安装或视觉模型识别能力不足，图片中文字可能无法被充分解析；本地小模型在复杂推理和长上下文总结方面也不如更大模型稳定。因此当前系统适合作为课程项目和轻量客服原型，后续还需要通过更多测试样例继续优化。"
    )
    page_break(doc)
    add_picture(doc, "chat.png", "图 1 桌面聊天界面：系统基于文档回答用户问题")
    add_picture(doc, "knowledge.png", "图 2 知识库界面：4 个文件已入库，共 122 个片段")
    add_picture(doc, "sources.png", "图 3 来源面板：展示命中文档与片段摘要")
    page_break(doc)

    doc.add_heading("7 遇到的问题与调整方案", level=1)
    table(
        doc,
        ["问题", "原因分析", "调整方案", "效果"],
        [
            ["本地大模型暂时不可用", "Ollama 服务未启动、模型未拉取或接口返回错误", "增加服务不可用提示，并在有检索结果时返回文档片段摘要", "系统不会直接中断，用户仍能看到相关依据"],
            ["图片内容无法有效解读", "普通文本解析无法读取图片中文字，视觉模型识别能力有限", "加入 OCR/视觉模型处理路径，支持图片格式上传", "具备图片资料处理能力，但质量依赖 OCR"],
            ["系统只会按文档回答", "原始 prompt 过度强调严格依据文档，缺少通用问题分流", "增加 general/document/hybrid/missing_evidence 路由", "可区分通用问题、文档问题和综合问题"],
            ["参数枚举回答漏项", "单次检索可能没有召回所有相关片段，模型也可能遗漏列表项", "增加多查询检索、必填项抽取和答案自查", "必选参数类问题完整性提升"],
            ["浏览器方式不方便使用", "原始 Gradio/FastAPI 需要通过浏览器访问", "改为 Tkinter 本地桌面应用", "交互更直接，符合本地程序使用习惯"],
        ],
        [1.4, 1.7, 1.8, 1.3],
    )
    doc.add_heading("7.1 文档图片识别问题", level=2)
    para(
        doc,
        "在测试过程中发现，文档中的图片内容并不一定能被系统有效理解。这是因为普通 PDF 或 Word 文本解析工具主要处理可复制文本，无法直接理解图片中的文字。如果接口参数表、流程图或说明文字以图片形式存在，就可能出现用户询问相关内容时，系统回答“文档中未找到相关信息”的情况。"
    )
    para(
        doc,
        "针对该问题，系统加入了 OCR 和视觉模型路径，并把图片格式纳入支持范围。这样上传图片或包含图片的资料时，系统会尝试把图片文字转成文本再入库。但该方案仍依赖 OCR 模型质量，因此后续应优先安装 RapidOCR 或配置 Tesseract，并用扫描件、截图、接口表格等资料做专项测试。"
    )
    doc.add_heading("7.2 通用问题与文档问题混合处理", level=2)
    para(
        doc,
        "另一个重要问题是：如果系统只要求模型严格依据文档回答，那么它会变成一个“只会查资料”的问答器，面对基础概念或常识问题时表现很差；但如果完全放开模型自由回答，又容易在业务问题上编造内容。"
    )
    para(
        doc,
        "因此本项目引入了问题路由机制。对于 general 类型问题，系统允许使用通用知识回答；对于 document 类型问题，系统必须依据检索片段；对于 hybrid 类型问题，系统可以把文档事实和通用解释结合起来；对于 missing_evidence 类型问题，系统明确说明缺少依据。这一设计让系统更接近实际客服：既能解释概念，也不会在缺少资料时假装知道。"
    )
    doc.add_heading("7.3 枚举类问题漏答问题", level=2)
    para(
        doc,
        "在接口文档问答中，用户经常询问“哪些参数是必选的”。这类问题要求系统列全字段，而不是只回答最相关的几个字段。测试中曾出现回答即梦 4.6 接口必选参数时漏掉 prompt 的情况，说明单次相似度检索和模型生成都可能造成漏项。"
    )
    para(
        doc,
        "调整方案是多层次的：首先对包含“必选、必填、required、参数”等关键词的问题扩大检索范围；其次增加多查询检索，把 prompt、Action、Version、req_key 等可能字段纳入查询；最后在回答后进行必填项抽取和自查。如果发现文档中有标注为必填但回答中没有出现的字段，就重新组织回答。"
    )
    page_break(doc)

    doc.add_heading("8 可行性验证", level=1)
    para(
        doc,
        "本项目重点验证了用户提出的智能体流程是否可落地。验证结果表明，该流程在当前项目中是可行的：上传阶段生成文档摘要和知识目录；提问阶段通过路由判断是否需要查文档；需要文档时再检索相关片段；回答后通过自查降低遗漏风险。该方案相比简单 RAG 更适合实际客服场景，因为它既能回答文档相关问题，也能处理部分通用概念问题，并能在缺少依据时提示用户补充资料。"
    )
    para(
        doc,
        "不过，方案质量仍取决于三个关键因素：第一是文档解析质量，尤其是图片和扫描件中的文字识别；第二是向量检索质量，完整 BGE-M3 通常优于 hash embedding 兜底方案；第三是本地大模型能力，小模型适合轻量问答，但复杂推理和长上下文综合仍有限。因此，后续优化应优先补齐 OCR、真实 embedding 权重和固定测试集。"
    )
    doc.add_heading("8.1 用户方案的核心思想", level=2)
    para(
        doc,
        "用户提出的方案可以概括为“先建立目录，再按需查文档，最后回答并自查”。该方案的优点是避免把全部文档片段无差别塞进 prompt，减少无关上下文对模型的干扰；同时也避免完全依赖模型记忆，让模型在需要业务事实时必须回到文档。"
    )
    doc.add_heading("8.2 与项目实现的对应关系", level=2)
    table(
        doc,
        ["用户设想", "项目中的实现", "验证方式"],
        [
            ["上传时生成总结性 Markdown 文档", "knowledge_index.py 生成文档摘要和目录", "上传文件后检查 data/knowledge_index"],
            ["提问时先让模型判断是否查文档", "question_router.py 输出问题模式和相关文件", "观察回答标签和 route 结果"],
            ["需要时再调用对应文档", "rag_chain.py 按 file_ids 检索相关片段", "查看来源面板是否命中文档"],
            ["结合文档、问题和上下文回答", "prompt 中包含历史对话、摘要和检索片段", "通过多轮问答测试上下文表现"],
            ["回答后进行一次自查", "answer_checker.py 和必填项抽取逻辑", "用必选参数问题验证漏项修正"],
        ],
        [1.9, 2.3, 2.1],
    )
    doc.add_heading("8.3 可行性结论", level=2)
    para(
        doc,
        "从当前实现看，该方案已经完成原型级验证。文档可以被入库，目录可以被建立，问题可以根据类型选择不同回答模式，答案也能显示来源。对于典型客服场景，该方案具有可行性，因为客服系统并不要求模型凭空掌握所有知识，而是要求模型能够找到依据、组织表达并说明来源。"
    )
    para(
        doc,
        "需要注意的是，可行性不等于完全成熟。当前系统仍然需要更强的 OCR、更稳定的 embedding、更系统的测试集和更完善的错误处理。只要这些基础能力继续补齐，TalkAgent 可以逐步从课程原型发展为更可靠的本地知识库客服。"
    )
    page_break(doc)

    doc.add_heading("9 核心代码节选", level=1)
    para(
        doc,
        "以下代码节选自项目核心逻辑，只保留能够说明系统工作方式的关键部分。完整代码位于项目 app 目录中，报告中不展开全部实现细节。"
    )
    doc.add_heading("9.1 回答模式提示词", level=2)
    para(
        doc,
        "回答模式提示词用于约束系统在不同问题类型下的行为：普通问题可以使用通用知识，文档问题必须依据资料，综合问题可以结合常识与文档，缺少依据时要明确说明。"
    )
    code_block(
        doc,
        '''
ANSWER_PROMPT = """你是一个可工作的 AI 客服助手。

你需要根据路由结果回答：
- general：直接使用通用知识回答，不要假装来自文档。
- document：必须依据参考文档和检索片段回答，不要编造文档外的业务事实。
- hybrid：文档事实以参考文档为准，解释、方案设计、步骤建议可以结合通用知识。
- missing_evidence：说明文档中未找到相关依据，不要编造具体业务事实。

回答开头必须使用其中一个标签：
【通用回答】、【基于文档】、【综合回答】、【缺少依据】
"""
''',
        "代码 1 回答模式约束",
    )

    doc.add_heading("9.2 检索增强流程", level=2)
    para(
        doc,
        "检索增强流程的作用是先根据问题路由判断是否需要查文档。如果需要，就调用向量检索；如果问题属于参数枚举类，还会扩大检索数量并增加扩展查询，减少漏掉关键字段的概率。"
    )
    code_block(
        doc,
        '''
def _retrieve_for_route(question: str, route: dict[str, Any]) -> list[Any]:
    if not route.get("needs_documents"):
        return []

    file_ids = route.get("relevant_file_ids") or None
    queries = _expanded_queries(question)
    docs = []
    seen = set()
    k = max(RETRIEVAL_K, 10) if _is_enumeration_question(question) else RETRIEVAL_K

    for query in queries:
        for doc in similarity_search(query, k=k, file_ids=file_ids):
            meta = doc.metadata or {}
            key = (meta.get("file_id", ""), meta.get("chunk_index", ""), doc.page_content[:80])
            if key not in seen:
                seen.add(key)
                docs.append(doc)
    return docs[: max(k, RETRIEVAL_K)]
''',
        "代码 2 根据路由执行文档检索",
    )

    doc.add_heading("9.3 最终问答链", level=2)
    para(
        doc,
        "最终问答链把历史对话、路由结果、文档摘要、检索片段和用户问题组合成 prompt，再交给本地模型。如果模型调用失败，则返回检索片段摘要作为降级结果。"
    )
    code_block(
        doc,
        '''
def answer_question(question: str, conversation_id: str | None = None) -> dict[str, Any]:
    history = chat_store.format_recent_history(conversation_id, CHAT_HISTORY_ROUNDS)
    route = route_question(question, history)
    docs = _retrieve_for_route(question, route)
    doc_context = _format_context(docs)
    doc_summaries = get_document_summaries(route.get("relevant_file_ids", []))

    prompt = f"""
历史对话：{history or "无"}
路由结果：{route}
相关文档摘要：{doc_summaries or "无"}
检索片段：{doc_context or "无"}
用户问题：{question}
请给出最终回答：
"""

    try:
        answer = generate_text(prompt, timeout=120)
        answer = check_answer(question, answer, route, doc_summaries + doc_context)
        answer = _normalize_answer_label(answer, route)
        answer = _enforce_required_items(question, answer, route, docs)
    except Exception as exc:
        answer = "本地大模型暂时不可用，先返回检索到的相关文档片段摘要。"

    return {"answer": answer, "sources": _sources_for_answer(answer, docs), "route": route}
''',
        "代码 3 问答链主流程",
    )
    page_break(doc)

    doc.add_heading("10 总结与展望", level=1)
    para(
        doc,
        "通过本次项目，实现了一个完整的本地 AI 客服原型。系统能够将上传文档转化为知识库，并通过 RAG 技术完成检索增强问答；同时加入问题路由、知识目录、答案自查和降级策略，使系统比单纯的文档问答更加接近实际客服工作流程。"
    )
    para(
        doc,
        "后续可以从三个方向继续完善：一是增强 OCR 和图片理解能力，让图片资料也能稳定进入知识库；二是补齐并优化 BGE-M3 embedding 模型，提高语义检索准确率；三是建立固定测试集，对常见业务问题、综合问题、无依据问题和异常情况进行持续评估。总体来看，TalkAgent 证明了在不重新训练大模型的情况下，通过 RAG 和智能体流程也可以构建一个可工作的本地 AI 客服系统。"
    )

    doc.add_heading("参考文献", level=1)
    for ref in [
        "[1] Lewis P., Perez E., Piktus A., et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. 2020.",
        "[2] Ollama 官方文档. 本地大语言模型运行与管理说明.",
        "[3] ChromaDB 官方文档. 向量数据库与相似度检索说明.",
        "[4] BAAI. BGE-M3 Embedding Model 技术说明.",
        "[5] Python 官方文档. Tkinter 图形界面库说明.",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(ref)
        set_run_font(r, 10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
