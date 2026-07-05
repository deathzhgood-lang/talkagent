from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "TalkAgent_PPT讲解脚本.docx"

FONT_CN = "宋体"
FONT_EN = "Times New Roman"
HEADING = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)


def set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def setup(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = FONT_EN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = HEADING
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def para(doc: Document, text: str, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r, 11)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, 10.5)


def label(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, 11, True, HEADING)


SLIDES = [
    {
        "title": "第 1 页：本地 AI 客服系统项目汇报",
        "summary": "封面页，展示项目名称、核心指标：已入库 4 个文件、122 个文本片段、运行形态为桌面端。",
        "script": [
            "各位老师好，我们小组展示的项目是 TalkAgent 本地 AI 客服系统。这个项目的目标是把用户上传的文档资料变成一个可以对话的知识库，让用户像问客服一样提问，系统能够根据资料生成回答，并展示回答来源。",
            "这页右侧的三个数字是当前项目的运行结果：系统已经完成 4 个文件的入库，生成了 122 个文本片段，并且最终运行形态是 Windows 桌面端应用。也就是说，它不是一个只能停留在代码里的实验，而是已经具备可交互界面的本地应用。",
        ],
        "qa": [
            "如果老师问“为什么要做本地 AI 客服”：可以回答，本地部署可以保护私有资料，不需要把内部文档上传到云端，同时也方便课程项目演示和离线测试。",
            "如果老师问“122 个文本片段是什么意思”：可以解释，长文档会被切成多个小片段，每个片段单独进入向量库，后续用户提问时系统会检索最相关的片段。",
        ],
    },
    {
        "title": "第 2 页：项目概述",
        "summary": "说明项目解决的问题：上传资料、建立知识库、自然语言提问、基于文档回答、显示来源。",
        "script": [
            "这一页主要说明项目要解决什么问题。传统客服系统通常依赖人工整理问答规则，维护成本比较高；而通用大模型虽然语言能力强，但它不知道我们上传的具体业务文档，直接回答时可能会编造。",
            "TalkAgent 的思路是：用户先上传接口文档、岗位资料或业务说明，系统自动把这些资料建立成本地知识库。用户提问时，系统先查找相关资料，再组织成客服式回答，并给出来源。这样既利用了大模型的表达能力，又让回答有文档依据。",
        ],
        "qa": [
            "“基于文档回答”指的是回答内容主要来自上传资料，而不是模型凭空生成。",
            "“来源展示”就是告诉用户答案参考了哪个文件、哪个片段，这样用户可以核对答案是否可靠。",
            "如果问“为什么要标注通用回答、基于文档、综合回答”：因为不同问题需要不同回答策略，标注后用户能知道答案依据来自哪里。",
        ],
    },
    {
        "title": "第 3 页：核心原理 RAG",
        "summary": "解释 RAG：上传文档、生成片段、向量入库、问题路由、生成与自查。",
        "script": [
            "这一页是项目的核心原理，叫 RAG，也就是检索增强生成。它不是重新训练一个模型，而是在用户提问时先从资料库中检索相关内容，再把这些内容交给大模型生成回答。",
            "流程可以分成五步：第一步上传文档，第二步把长文档切成片段，第三步把片段转换成向量并存入 ChromaDB，第四步用户提问时检索相关片段，第五步让大模型结合片段生成回答，并进行自查。",
        ],
        "qa": [
            "RAG 的英文是 Retrieval-Augmented Generation，意思是“检索增强生成”。检索负责找资料，生成负责把资料组织成自然语言答案。",
            "为什么不直接训练模型：训练成本高、数据准备复杂，而且课程项目环境有限。RAG 可以不训练模型，只要更新知识库就能更新知识。",
            "向量是什么：可以理解为文本的数字化语义表示，语义相近的文本在向量空间中距离更近。",
        ],
    },
    {
        "title": "第 4 页：技术栈",
        "summary": "介绍 Tkinter、Ollama、BGE-M3/hash embedding、ChromaDB、文档解析、智能体流程。",
        "script": [
            "这一页介绍项目用到的主要技术。界面层使用 Tkinter，因为它适合快速构建 Windows 本地桌面应用。模型服务使用 Ollama，本地模型名称是 codex-app，用于生成最终回答。",
            "向量检索部分默认使用 BGE-M3 embedding 模型，把问题和文档片段转换为向量；如果权重不完整，就使用 hash embedding 兜底。向量数据库使用 ChromaDB，负责保存文档片段和向量索引。文档解析使用 PyMuPDF、docx2txt 和 OCR/视觉模型。智能体流程由 Router、RAG 和 Checker 组成，分别负责判断、检索、回答和自查。",
        ],
        "qa": [
            "Tkinter 是 Python 自带的图形界面库，优点是简单、无需浏览器，适合本地桌面程序。",
            "Ollama 是本地运行大模型的工具，可以在本机启动模型服务，通过接口调用模型。",
            "ChromaDB 是向量数据库，用来保存向量并做相似度搜索。",
            "BGE-M3 是 embedding 模型，用于把文本转换成语义向量；hash embedding 是兜底方案，能跑通流程，但语义效果不如真实 embedding。",
        ],
    },
    {
        "title": "第 5 页：你的关键想法",
        "summary": "重点解释用户提出的智能体流程：先看目录，再决定是否查文档。",
        "script": [
            "这一页是项目中比较关键的设计思路。最初普通 RAG 的做法是用户一提问就直接检索文档，但这样可能会遇到两个问题：一是所有问题都去查文档，普通常识问题也被限制住；二是如果文档很多，检索范围太大，可能找到不相关内容。",
            "所以这里采用的想法是：上传文档时先生成 Markdown 总结和检索目录；用户提问时，系统先把问题和目录进行判断，决定是否需要查文档、查哪些文档。如果需要，再检索相关片段；最终回答后再做一次遗漏和事实自查。这样流程更像一个真正的智能体，而不是简单的搜索器。",
        ],
        "qa": [
            "“目录”不是文件目录，而是系统对已上传文档内容的摘要和索引，帮助模型知道知识库里有什么。",
            "“先判断再检索”的好处是减少无关检索，提高回答针对性，也允许系统回答普通概念问题。",
            "如果问“这个流程和普通 RAG 有什么区别”：普通 RAG 往往直接检索；这里多了知识目录、问题路由和答案自查。",
        ],
    },
    {
        "title": "第 6 页：遇到的问题",
        "summary": "说明开发中遇到的四个问题及调整方案：Ollama 不可用、图片识别、只按文档答、必选参数漏答。",
        "script": [
            "这一页总结项目过程中遇到的主要问题。第一个问题是 Ollama 服务不可用时回答会失败，所以我们加入了检索片段摘要兜底，让模型掉线时系统仍能返回有用信息。",
            "第二个问题是图片里的文字没有被读到，所以增加了 OCR 和 Vision 路径，并支持图片上传。第三个问题是系统最开始只会根据文档回答，常识问题不会答，所以加入了 general、document、hybrid 路由。第四个问题是必选参数类问题容易漏项，比如漏掉 prompt，所以加入多查询检索、必填项抽取和答案自查。",
        ],
        "qa": [
            "OCR 是 Optical Character Recognition，中文叫光学字符识别，用来把图片中的文字转成可检索文本。",
            "Vision 路径指的是用视觉模型理解图片内容，但小模型识别能力有限，所以 OCR 更适合文字识别。",
            "general/document/hybrid 分别表示通用回答、文档回答和综合回答。",
            "必选参数漏答本质是召回和生成两个环节都可能丢信息，所以需要扩展检索和自查补偿。",
        ],
    },
    {
        "title": "第 7 页：方案落地",
        "summary": "说明智能体流程对应到具体模块：knowledge_index、question_router、rag_chain、answer_checker、desktop。",
        "script": [
            "这一页说明前面的设计并不是停留在想法层，而是已经拆成了具体模块。knowledge_index.py 负责上传后生成文档摘要和目录；question_router.py 负责判断问题是否需要文档；rag_chain.py 负责执行检索和回答；answer_checker.py 负责检查遗漏和依据不足；desktop.py 负责把结果展示给用户。",
            "这样的模块拆分有一个好处：每个模块职责比较清楚，后续如果要优化 OCR、替换模型或者调整路由策略，不需要重写整个系统。"
        ],
        "qa": [
            "knowledge_index.py 可以理解为知识库目录生成器。",
            "question_router.py 是问题分流器，决定问题走通用回答、文档回答还是综合回答。",
            "rag_chain.py 是问答主链路，把历史、摘要、检索片段和用户问题组合起来。",
            "answer_checker.py 是最后的质量检查，用来降低漏项和事实不一致。",
        ],
    },
    {
        "title": "第 8 页：可行性验证 A",
        "summary": "用知识库截图验证文档可以入库，目录可以建立。",
        "script": [
            "这一页是第一部分可行性验证，验证的是知识库是否真的建立起来。截图中可以看到系统已经入库了 4 个文件，并生成了 122 个文本片段。",
            "这个验证很重要，因为如果文档没有成功入库，后面的问答就没有依据。现在知识库界面能显示文件名、片段数和 file_id，说明上传、解析、分块和向量入库链路已经跑通。"
        ],
        "qa": [
            "file_id 是系统给每个上传文件生成的唯一标识，用来区分不同文档。",
            "片段数不是文件数，一个文件会被切成多个片段，方便精确检索。",
            "如果老师问“为什么要分块”：因为大模型上下文有限，长文档整体检索不精确，分块可以提高召回效果。",
        ],
    },
    {
        "title": "第 9 页：可行性验证 B",
        "summary": "用聊天截图验证系统能基于资料回答问题。",
        "script": [
            "这一页展示第二部分验证：系统能不能基于资料回答问题。测试问题是“如果我要去应聘 AI 应用开发，需要怎样才能被录用”。这个问题不是简单查一个关键词，而是需要从文档中整理岗位能力要求。",
            "从截图可以看到，系统回答中包含技术能力、产品思维、伦理合规、协作能力和实践经验等内容，说明它不是只做关键词匹配，而是把检索到的资料组织成了结构化回答。"
        ],
        "qa": [
            "这里的“基于文档”表示回答主要来自上传的岗位资料，而不是模型自由发挥。",
            "如果问“模型有没有可能编造”：有可能，所以系统通过来源展示、回答标签和答案自查来降低风险，但不能完全替代人工审核。",
            "结构化回答是指把文档中的零散信息整理成条目，方便用户阅读。",
        ],
    },
    {
        "title": "第 10 页：可行性验证 C",
        "summary": "用来源面板截图验证回答可以追溯。",
        "script": [
            "这一页验证的是来源追溯能力。一个 AI 客服如果只给答案但不给依据，用户很难判断它是否可靠。这里系统展示了命中的文件名、片段编号和原文摘要。",
            "这说明整个链路形成了闭环：先建立知识库，再路由问题，再检索文档，再生成回答，最后把来源展示出来。这个结果证明“先看目录、再决定查文档、再回答和自查”的方案在当前项目中是可落地的。"
        ],
        "qa": [
            "来源追溯的作用是提高可解释性，让用户能核对答案依据。",
            "片段编号 chunk_index 表示答案引用的是文件中的哪一段。",
            "如果老师问“来源能不能保证答案一定正确”：不能百分百保证，但它能让用户检查依据，是降低幻觉的重要手段。",
        ],
    },
    {
        "title": "第 11 页：结论",
        "summary": "总结项目可行，但质量取决于文档解析、检索和模型能力。",
        "script": [
            "最后一页是结论。当前项目已经证明：在不重新训练模型的情况下，也可以通过 RAG 和智能体流程做出一个可工作的本地 AI 客服。它已经跑通了文档上传、知识库入库、检索问答、来源追溯和桌面交互。",
            "但同时也要说明，系统质量仍然取决于三个方面：第一是文档解析能力，尤其是图片和扫描件；第二是检索能力，真实 BGE-M3 embedding 会比 hash 兜底更好；第三是本地模型能力，小模型适合轻量问答，但复杂推理仍有限。后续优化方向是增强 OCR、补齐 BGE-M3 权重，并建立固定测试集持续评估。"
        ],
        "qa": [
            "如果老师问“项目最大的创新点”：可以回答，是在普通 RAG 基础上加入了知识目录、问题路由和答案自查，让系统更接近实际客服流程。",
            "如果老师问“项目不足”：可以回答，图片识别、embedding 质量和小模型能力仍是限制，需要后续继续优化。",
            "如果老师问“能否实际使用”：可以回答，当前已达到原型可用，适合小规模本地知识库问答；正式生产环境还需要更多测试和稳定性优化。",
        ],
    },
]


def build():
    doc = Document()
    setup(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TalkAgent 项目 PPT 讲解脚本")
    set_font(r, 18, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("按 PPT 页码顺序整理，可作为答辩旁白和老师追问时的解释准备")
    set_font(r, 10.5, False, MUTED)

    doc.add_heading("使用说明", level=1)
    para(doc, "本文档按 PPT 页码顺序编写。每一页包含页面内容概括、可直接讲述的旁白，以及老师可能追问时可以使用的概念解释。答辩时可以优先讲“讲解旁白”，如果老师追问具体概念，再参考“可能追问与解释”。")

    for slide in SLIDES:
        doc.add_heading(slide["title"], level=1)
        label(doc, "页面内容概括")
        para(doc, slide["summary"])
        label(doc, "讲解旁白")
        for text in slide["script"]:
            para(doc, text)
        label(doc, "可能追问与概念解释")
        for item in slide["qa"]:
            bullet(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
